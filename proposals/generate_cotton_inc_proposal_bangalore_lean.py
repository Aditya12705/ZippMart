"""Generate a lean Bangalore-vendor-only Cotton Incorporated pre-proposal."""

from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = r"d:\Checkout System\proposals\Cotton_Incorporated_RFP_2027_PreProposal_Bangalore_Lean_Budget.docx"


def shade(cell, color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x5D, 0x3A)
    return h


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        shade(table.rows[0].cells[i], "D9EAD3")
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = str(value)
            for p in table.rows[r].cells[c].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for c, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[c].width = Cm(width)
    return table


def fmt_inr(value):
    return f"₹{value:,.2f}"


BUDGET_ITEMS = [
    {
        "sl": 1,
        "item": "Cotton linters for cotton reinforcement / partial matrix replacement",
        "qty": 30,
        "unit": "kg",
        "unit_price": 25.00,
        "amount": 750.00,
        "vendor": "Vicram Yarns And Fibers",
        "address": "303, 2nd Floor, Siri Aroma Apartments, 90 Kempanna Layout, Gowdanapalya, Uttrahalli, Bengaluru, Karnataka 560061",
        "basis": "TradeIndia listing: Cotton Linters - Fiber Length Max 6 mm at ₹25/kg",
    },
    {
        "sl": 2,
        "item": "Food-grade maize starch for baseline starch-pectin film trials",
        "qty": 25,
        "unit": "kg",
        "unit_price": 39.00,
        "amount": 975.00,
        "vendor": "Well Thought Chemicals",
        "address": "355, 9th Cross, 4th Main Road, 4th Phase, Peenya Industrial Area, Near NTTF Circle, Bengaluru, Karnataka 560058",
        "basis": "IndiaMART listing: White Maize Starch Powder at ₹39/kg in Bengaluru",
    },
    {
        "sl": 3,
        "item": "Glycerine (>99%, laboratory use) as plasticizer",
        "qty": 5,
        "unit": "kg",
        "unit_price": 250.00,
        "amount": 1250.00,
        "vendor": "Paxal Chemical Industry Private Limited",
        "address": "Municipal No. 6, Railway Parallel Road, Nehru Nagar, Sheshadripuram, Bengaluru, Karnataka 560020",
        "basis": "IndiaMART listing: Glycerine for laboratory use at ₹250/kg",
    },
    {
        "sl": 4,
        "item": "Glacial acetic acid for controlled acidification during film casting",
        "qty": 10,
        "unit": "kg",
        "unit_price": 100.00,
        "amount": 1000.00,
        "vendor": "Miracle Ingredients LLP",
        "address": "No. 4, 2nd Cross, Sri Ranganagara Outer Ring Road, Pantharapalya, 100 Ft Ring Road, Bengaluru, Karnataka 560039",
        "basis": "IndiaMART sitemap/listing: Glacial Acetic Acid approx. ₹100/kg",
    },
    {
        "sl": 5,
        "item": "Borosilicate beakers for dedicated casting and formulation work",
        "qty": 2,
        "unit": "piece",
        "unit_price": 500.00,
        "amount": 1000.00,
        "vendor": "Nandini Marketing Company",
        "address": "No. 174, Nandini Mansion, 3rd Main, 5th Cross, BHEL Layout, Kenchenhalli, Rajarajeshwari Nagar, Bengaluru, Karnataka 560098",
        "basis": "Supplier listing: BORO+ Borosilicate Beaker at ₹500/piece",
    },
]


TOTAL_BUDGET = round(sum(item["amount"] for item in BUDGET_ITEMS), 2)


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRE-PROPOSAL FOR RESEARCH FUNDING")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x2E, 0x5D, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Request for Proposals 2027\nProduct Development and Implementation Division\nCotton Incorporated")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Low-Cost Development of Cotton-Linter Reinforced\n"
        "Biodegradable Starch-Pectin Films for Single-Use Packaging"
    )
    r.italic = True
    r.font.size = Pt(13)

    doc.add_paragraph()
    cover_lines = [
        ("Principal Investigator:", "[Dr. _________________________]"),
        ("Department:", "[Department of _________________________]"),
        ("Institution:", "[University / Institute Name, Bengaluru / India]"),
        ("Target RFP Category:", "Priority Area 3a - Cotton-derived films, wraps, or barrier coatings"),
        ("Project Duration:", "6 months (proof-of-concept stage)"),
        ("Budget Request:", fmt_inr(TOTAL_BUDGET)),
        ("Procurement Policy:", "All priced vendors in this draft are Bengaluru-based"),
        ("Document Date:", date.today().strftime("%d %B %Y")),
    ]
    for label, value in cover_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label} ").bold = True
        p.add_run(value)

    doc.add_page_break()

    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "This pre-proposal presents a deliberately low-cost, short-duration proof-of-concept "
        "study for Cotton Incorporated. The proposed work builds directly on our prior "
        "biodegradable plastic formulation based on cornstarch, pectin, water, vinegar/acetic "
        "acid, and glycerin, and extends it by introducing cotton linters as a cotton-derived "
        "reinforcement and partial matrix replacement. Rather than requesting a large capital "
        "budget, this version is designed around materials only. Existing departmental equipment, "
        "student effort, pectin stock from previous work, and routine analytical facilities will be "
        "used as in-kind support."
    )
    doc.add_paragraph(
        "The goal is to produce an initial set of cotton-linter reinforced biodegradable films for "
        "single-use packaging and demonstrate whether cotton addition improves strength, handling, "
        "and biodegradation behavior relative to the earlier starch-based formulation. This makes "
        "the proposal practical, focused, and aligned with Cotton Incorporated's preference for "
        "well-defined research capable of generating preliminary outcomes in 6-12 months."
    )

    add_heading(doc, "2. Alignment with the RFP", 1)
    add_table(
        doc,
        ["RFP Priority", "How this proposal fits"],
        [
            [
                "Priority Area 3a - Cotton-derived biodegradable alternatives to polyethylene",
                "Cotton linters are introduced directly into biodegradable film development for packaging applications.",
            ],
            [
                "Priority Area 5 - Valorization of cotton processing byproducts",
                "Cotton linters, a low-value cotton byproduct, are converted into a packaging material component.",
            ],
            [
                "Targeted, short-duration research",
                "The work is scoped as a 6-month proof-of-concept using existing institutional facilities.",
            ],
        ],
        col_widths=[6, 10],
    )

    add_heading(doc, "3. Background and Concept", 1)
    doc.add_paragraph(
        "Our laboratory has already prepared biodegradable plastic sheets using cornstarch, pectin, "
        "water, vinegar, and glycerin. Those earlier trials established a workable film-forming "
        "process but were not yet centered on cotton. The present proposal modifies that successful "
        "base formulation by incorporating micronized cotton linters obtained locally in Bengaluru. "
        "This allows the research to remain low-cost while still addressing Cotton Incorporated's "
        "core interest: expansion of cotton into biodegradable single-use material applications."
    )
    doc.add_paragraph(
        "At this stage, the project does not aim to purchase expensive testing equipment, hire new "
        "staff, or outsource certification-scale biodegradation studies. Instead, it focuses on the "
        "essential first question: can a cotton-linter reinforced starch-pectin film be produced "
        "reliably at lab scale with acceptable flexibility, handling, and observable biodegradation?"
    )

    add_heading(doc, "4. Objectives", 1)
    objectives = [
        "To prepare cotton-linter reinforced biodegradable films using the existing starch-pectin-glycerin-acid formulation platform.",
        "To compare films containing 0%, 10%, 20%, and 30% cotton linter loading with respect to film formation, thickness, and handling quality.",
        "To perform initial mechanical and biodegradation observations using existing institutional facilities.",
        "To identify one optimized low-cost formulation for follow-on full proposal development.",
    ]
    for objective in objectives:
        doc.add_paragraph(objective, style="List Bullet")

    add_heading(doc, "5. Proposed Methodology", 1)
    method_steps = [
        "Micronize or manually refine cotton linters to improve dispersion in the polymer slurry.",
        "Prepare the baseline starch-pectin formulation using prior laboratory protocol.",
        "Introduce cotton linters at increasing proportions and cast films on flat glass / beaker-supported casting surfaces.",
        "Dry films under controlled room-temperature and low-temperature oven conditions available in the department.",
        "Record film thickness, visual uniformity, flexibility, cracking behavior, and peelability.",
        "Use existing departmental instruments for simple tensile comparison if available; otherwise document fold endurance and handling strength.",
        "Conduct preliminary in-house biodegradation screening through compost/soil burial and periodic mass-loss observation.",
    ]
    for step in method_steps:
        doc.add_paragraph(step, style="List Number")

    add_heading(doc, "6. Six-Month Work Plan", 1)
    add_table(
        doc,
        ["Month", "Activity", "Output"],
        [
            ["1", "Procurement of Bengaluru-sourced materials and replication of prior starch film", "Baseline formulation reproduced"],
            ["2", "Cotton linter preparation and dispersion trials", "Feasible cotton incorporation range identified"],
            ["3", "Casting of four main formulations", "Film samples prepared"],
            ["4", "Film comparison and optimization", "One best formulation shortlisted"],
            ["5", "Initial biodegradation screening and mechanical observations", "Preliminary performance dataset"],
            ["6", "Data consolidation and full-proposal preparation", "Optimized formulation and next-stage plan"],
        ],
        col_widths=[2, 7, 7],
    )

    add_heading(doc, "7. Expected Preliminary Outcomes", 1)
    outcomes = [
        "A reproducible cotton-linter reinforced biodegradable film formulation.",
        "Preliminary evidence that cotton byproducts can be integrated into the existing starch-pectin film system.",
        "A low-cost materials dataset suitable for a full proposal to Cotton Incorporated.",
        "A Bengaluru-localized sourcing model demonstrating practical procurement within India.",
    ]
    for outcome in outcomes:
        doc.add_paragraph(outcome, style="List Bullet")

    add_heading(doc, "8. Budget Strategy", 1)
    doc.add_paragraph(
        "This revised budget is intentionally minimal. No new equipment, travel, publication fees, "
        "or dedicated project salaries are requested. Pectin from prior bioplastic work, distilled "
        "water, common laboratory glassware, hotplate/stirrer access, oven access, and routine "
        "student effort will be provided as in-kind support by the host department."
    )
    doc.add_paragraph(
        "Only core materials that must be freshly procured are included below. All priced vendors "
        "are based in Bengaluru, Karnataka, India."
    )

    add_heading(doc, "9. Detailed Budget - Bengaluru Vendors Only", 1)
    budget_rows = []
    for item in BUDGET_ITEMS:
        budget_rows.append(
            [
                item["sl"],
                item["item"],
                f'{item["qty"]} {item["unit"]}',
                fmt_inr(item["unit_price"]),
                item["vendor"],
                fmt_inr(item["amount"]),
            ]
        )
    add_table(
        doc,
        ["Sl.", "Item", "Qty", "Unit Price", "Bengaluru Vendor", "Amount"],
        budget_rows,
        col_widths=[0.8, 6.5, 1.5, 1.8, 4.3, 1.8],
    )

    p = doc.add_paragraph()
    p.add_run("Total Budget Requested: ").bold = True
    total_run = p.add_run(fmt_inr(TOTAL_BUDGET))
    total_run.bold = True
    total_run.font.size = Pt(12)
    total_run.font.color.rgb = RGBColor(0x2E, 0x5D, 0x3A)

    add_heading(doc, "10. Item-wise Vendor Details", 1)
    vendor_rows = []
    for item in BUDGET_ITEMS:
        vendor_rows.append(
            [
                item["vendor"],
                item["address"],
                item["basis"],
                fmt_inr(item["amount"]),
            ]
        )
    add_table(
        doc,
        ["Vendor", "Address", "Quoted Price Basis", "Allocated Amount"],
        vendor_rows,
        col_widths=[3.3, 6.0, 4.0, 2.0],
    )

    add_heading(doc, "11. In-Kind Institutional Support (Not Charged to the Project)", 1)
    in_kind = [
        "Existing pectin stock from previous biodegradable plastic work",
        "Distilled water and basic laboratory reagents already available in the department",
        "Hotplate/stirrer, drying oven, balance, desiccator, and basic glassware already available",
        "Faculty supervision and existing student effort",
        "Routine internal testing support through departmental / institutional facilities",
    ]
    for item in in_kind:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "12. Why This Budget Is Realistic", 1)
    doc.add_paragraph(
        "The earlier draft looked more like a full-scale research grant with capital equipment, "
        "travel, manpower, outsourced testing, and publication charges. That scale is not necessary "
        "for a first-stage Cotton Incorporated pre-proposal. This revised version asks only for the "
        "minimum material support required to generate convincing preliminary data. If the proof-of-"
        "concept succeeds, a later full proposal can justify expanded testing, certification, and "
        "scale-up. For the pre-proposal stage, a lean budget is more credible and easier to defend."
    )

    add_heading(doc, "13. Conclusion", 1)
    doc.add_paragraph(
        "This revised pre-proposal offers a low-risk, low-cost, Bengaluru-sourced pathway to test "
        "whether cotton linters can enhance an already established biodegradable plastic formulation. "
        "The requested budget is intentionally kept to the smallest practical level while still "
        "making the study executable. The outcome will be a strong, data-backed basis for a fuller "
        "proposal if Cotton Incorporated expresses interest."
    )

    doc.add_paragraph()
    doc.add_paragraph("Principal Investigator Signature: _________________________    Date: _______________")
    doc.add_paragraph("Name: [Dr. _________________________]")
    doc.add_paragraph("Head of Department Endorsement: _________________________    Date: _______________")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Generated: {path}")
    print(f"Lean budget total: INR {TOTAL_BUDGET:,.2f}")
