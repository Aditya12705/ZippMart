"""Generate Cotton Incorporated RFP 2027 Pre-Proposal DOCX."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

OUTPUT = r"d:\Checkout System\proposals\Cotton_Incorporated_RFP_2027_PreProposal_Biodegradable_Cotton_Films.docx"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x5D, 0x3A)
    return h


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "D9EAD3")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def fmt_inr(amount):
    return f"₹{amount:,.2f}"


# --- Budget line items with vendor details ---
# All unit prices sourced from publicly listed vendor catalogues (May 2026).
# GST @ 18% applied where vendor prices exclude tax.

GST_RATE = 0.18
NON_TAXABLE_CATEGORIES = {"Manpower"}


def line_amount(qty, unit_price):
    return round(qty * unit_price, 2)


def compute_budget_items():
    raw = [
        (1, "Raw Materials", "Bleached Cotton Linter Pulp (medium viscosity, food/paper grade)", 80, "kg", 130.00,
         "Moirai Cotton Pulp Private Limited", "Salemgarh, Hisar, Haryana – 125052"),
        (2, "Raw Materials", "Pectin Powder (Food/Pharma Grade, CAS 9000-69-5, >99% purity)", 10, "kg", 1098.00,
         "Amit Hydrocolloids", "121, Raja Industrial Estate, PK Road, Mumbai – 400080"),
        (3, "Raw Materials", "Glycerol GR 99%+ (AR Grade, CAS 56-81-5)", 15, "L", 1401.80,
         "Otto Kemi (OTTO Chemie Pvt. Ltd.)", "Mumbai, Maharashtra"),
        (4, "Raw Materials", "Maize Starch Food Grade (CAS 9005-25-8, 99% purity) – comparative baseline", 25, "kg", 35.00,
         "Triveni Chemicals India", "Vapi, Gujarat"),
        (5, "Raw Materials", "Acetic Acid Glacial AR (99.7%, CAS 64-19-7)", 5, "L", 488.00,
         "Alpha Chemika", "102, Savagan Heights, Andheri West, Mumbai – 400058"),
        (6, "Raw Materials", "Sulphuric Acid AR (64–98%, CAS 7664-93-9) – for CNC hydrolysis", 2.5, "L", 419.80,
         "Labitems India", "Online laboratory supplier, PAN India dispatch"),
        (7, "Raw Materials", "Sodium Hydroxide Pellets AR (for neutralization & pH control)", 1, "kg", 285.00,
         "HiMedia Laboratories Pvt. Ltd.", "MIDC Wagle Estate, Thane – 400604"),
        (8, "Raw Materials", "Food-grade White Vinegar (4–5% acetic acid) – formulation trials", 20, "L", 45.00,
         "Local FSSAI-certified food supplier (proforma)", "Institution city – local market"),
        (9, "Consumables", "Borosilicate glass beakers (100–1000 mL assorted, Borosil brand)", 1, "set", 4200.00,
         "Borosil Scientific Ltd. / authorized dealer", "Worli, Mumbai / regional distributor"),
        (10, "Consumables", "Teflon-coated film casting plates (300×300 mm) + casting knife set", 1, "set", 8500.00,
         "Royal Scientific Solutions", "Delhi – IndiaMART verified supplier"),
        (11, "Consumables", "Whatman Filter Paper Grade 1 (125 mm, pack of 100)", 5, "pack", 680.00,
         "Cytiva / authorized India distributor", "Bengaluru, Karnataka"),
        (12, "Consumables", "Aluminium foil, HDPE storage bags, desiccant sachets, lab labels", 1, "lot", 3200.00,
         "Local scientific consumables vendor (proforma)", "Institution city"),
        (13, "Equipment", "Magnetic Hotplate Stirrer (5 L capacity, digital temp. control, 350°C max)", 2, "unit", 12500.00,
         "Labindia Analytical Instruments Pvt. Ltd.", "Mumbai, Maharashtra"),
        (14, "Equipment", "Vacuum Oven (25 L, max 250°C, digital controller)", 1, "unit", 78000.00,
         "Royal Scientific / EIE Instruments (comparable model quote)", "Ahmedabad, Gujarat"),
        (15, "Equipment", "Digital Micrometer Thickness Gauge (0–10 mm, ±0.001 mm)", 1, "unit", 9800.00,
         "Mitutoyo India / Insize authorized dealer", "Pune, Maharashtra"),
        (16, "Equipment", "Portable pH Meter with ATC (0.01 pH resolution, calibrated)", 1, "unit", 6500.00,
         "Labindia Analytical Instruments Pvt. Ltd.", "Mumbai, Maharashtra"),
        (17, "Equipment", "Universal Testing Machine, 10 kN capacity (tensile mode with film grips)", 1, "unit", 470000.00,
         "Eqvimech Private Limited", "Jekegram, Thane – 400606"),
        (18, "Testing Services", "Compost Biodegradation Testing – ISO 14855-1 / ASTM D5338 (per formulation)", 3, "sample", 35000.00,
         "Akshar Global Exports (ISO/IEC 17025 accredited partner lab)", "Kolkata, West Bengal"),
        (19, "Testing Services", "Disintegration & Ecotoxicity screening – EN 13432 supplementary panel", 1, "batch", 42000.00,
         "Akshar Global Exports", "Kolkata, West Bengal"),
        (20, "Testing Services", "Marine environment simulation – water column & sediment biodegradation study", 1, "study", 55000.00,
         "Environmental Research & Remediation Laboratory (ERRL)", "India – ASTM D6691 aligned protocol"),
        (21, "Testing Services", "DSC & TGA thermal analysis (per sample, outsourced NABL lab)", 8, "sample", 2200.00,
         "J. K. Analytical Laboratory & Research Centre", "Ahmedabad, Gujarat"),
        (22, "Testing Services", "Water Vapor Transmission Rate (WVTR) – ASTM E96 gravimetric method", 6, "sample", 4500.00,
         "J. K. Analytical Laboratory & Research Centre", "Ahmedabad, Gujarat"),
        (23, "Manpower", "Junior Research Fellow (Project) – 18 months @ ₹37,000/month (UGC JRF rate 2025–26)", 18, "month", 37000.00,
         "As per UGC/CSIR fellowship norms", "University Grants Commission, New Delhi"),
        (24, "Manpower", "Project Assistant (Technical) – 12 months @ ₹22,000/month", 12, "month", 22000.00,
         "Institutional project staff emoluments", "Host institution payroll"),
        (25, "Manpower", "Principal Investigator Research Allowance – 18 months @ ₹12,000/month", 18, "month", 12000.00,
         "Institutional PI honorarium norms", "Host institution"),
        (26, "Travel", "Domestic travel – vendor visits, testing lab coordination (4 trips)", 4, "trip", 8500.00,
         "IRCTC / domestic airlines (proforma)", "India"),
        (27, "Travel", "International travel – Cotton Incorporated progress review meeting, USA (1 PI + 1 JRF)", 1, "visit", 285000.00,
         "Institutional travel desk (economy airfare + per diem estimate)", "USA"),
        (28, "Publication", "Open-access journal publication charges (1 article)", 1, "article", 45000.00,
         "Target journal APC (estimated)", "International publisher"),
    ]
    items = []
    taxable_subtotal = 0.0
    for row in raw:
        sl, cat, desc, qty, unit, unit_price, vendor, loc = row
        excl = line_amount(qty, unit_price)
        taxable = cat not in NON_TAXABLE_CATEGORIES
        gst = round(excl * GST_RATE, 2) if taxable else 0.0
        if cat in ("Raw Materials", "Consumables", "Equipment", "Testing Services"):
            taxable_subtotal += excl
        items.append((sl, cat, desc, qty, unit, unit_price, vendor, loc, gst, excl))
    contingency_excl = round(taxable_subtotal * 0.05, 2)
    contingency_gst = round(contingency_excl * GST_RATE, 2)
    items.append((29, "Contingency", "Contingency @ 5% on Items 1–22 (materials, equipment, testing)", 1, "lot",
                  contingency_excl, "—", "—", contingency_gst, contingency_excl))
    return items


BUDGET_ITEMS = compute_budget_items()

def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # --- COVER PAGE ---
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PRE-PROPOSAL FOR RESEARCH FUNDING")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x2E, 0x5D, 0x3A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Request for Proposals 2027\nProduct Development and Implementation Division\nCotton Incorporated")
    r2.font.size = Pt(14)
    r2.bold = True

    doc.add_paragraph()
    proj = doc.add_paragraph()
    proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = proj.add_run(
        "Development of Cotton-Derived Biodegradable Films as\n"
        "Sustainable Alternatives to Single-Use Polyethylene Packaging"
    )
    r3.font.size = Pt(13)
    r3.italic = True

    doc.add_paragraph()
    fields = [
        ("Principal Investigator:", "[Dr. _________________________]"),
        ("Designation:", "[Professor / Associate Professor]"),
        ("Department:", "[Department of _________________________]"),
        ("Institution:", "[University / Institute Name, City, India]"),
        ("Email:", "[pi.email@institution.ac.in]"),
        ("Phone:", "[+91-_________________]"),
        ("Target Priority Area:", "Area 3 – Biodegradable Alternatives to Single-Use Plastics (Sub-area 3a)"),
        ("Project Duration:", "June 2026 – December 2027 (19 months)"),
        ("Pre-Proposal Submission Deadline:", "Sunday, 31 May 2026"),
        ("Document Date:", date.today().strftime("%d %B %Y")),
    ]
    for label, val in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label} ").bold = True
        p.add_run(val)

    doc.add_page_break()

    # --- TABLE OF CONTENTS placeholder ---
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Executive Summary",
        "2. Alignment with Cotton Incorporated Strategic Priorities",
        "3. Background and Prior Work",
        "4. Problem Statement and Research Gap",
        "5. Project Objectives and Hypotheses",
        "6. Methodology and Work Plan",
        "7. Timeline and Milestones",
        "8. Expected Outcomes and Deliverables",
        "9. Commercialization and Market Impact Potential",
        "10. Investigator Qualifications and Institutional Resources",
        "11. Detailed Budget with Vendor Quotations (INR)",
        "12. Budget Summary",
        "13. References",
        "Appendix A – Vendor Quotation Summary Sheet",
        "Appendix B – Pre-Proposal Submission Checklist",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # --- 1. EXECUTIVE SUMMARY ---
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "This pre-proposal requests funding from Cotton Incorporated to develop cotton-derived "
        "biodegradable films as commercially viable alternatives to single-use polyethylene (PE) "
        "in flexible packaging applications. Building on our laboratory's established expertise in "
        "polysaccharide-based biodegradable plastics—formulated from cornstarch, pectin, glycerin, "
        "and dilute acetic acid (vinegar)—we propose to systematically replace the starch matrix "
        "with bleached cotton linter pulp and cotton cellulose nanocrystals (CNCs), thereby "
        "directly valorizing cotton processing byproducts while delivering a product aligned with "
        "regulatory and consumer demand for plastic-free packaging."
    )
    doc.add_paragraph(
        "The project targets Priority Area 3a of the 2027 RFP: development of cotton-derived films, "
        "wraps, or barrier coatings as biodegradable alternatives to polyethylene. Within 6–12 months, "
        "we will deliver optimized film prototypes with documented mechanical properties, barrier "
        "performance, and validated biodegradation data under composting and marine-simulation "
        "conditions. All activities will be completed by December 2027."
    )

    # --- 2. ALIGNMENT ---
    add_heading(doc, "2. Alignment with Cotton Incorporated Strategic Priorities", 1)
    add_heading(doc, "2.1 Primary Priority Area", 2)
    doc.add_paragraph(
        "Priority Area 3 – Biodegradable Alternatives to Single-Use Plastics\n"
        "Sub-area 3a: Research the development of cotton-derived films, wraps, or barrier coatings "
        "as biodegradable alternatives to polyethylene in flexible packaging applications."
    )
    add_heading(doc, "2.2 Secondary Alignment", 2)
    doc.add_paragraph(
        "Priority Area 5 – Valorization of Cotton Processing Byproducts: Cotton linter pulp sourced "
        "from ginning/spinning waste streams serves as the primary feedstock, converting "
        "underutilized byproducts into high-value packaging materials."
    )
    add_table(doc,
        ["RFP Requirement", "Project Response"],
        [
            ["Applied research with near-term implementation pathway", "Film casting protocol scalable to pilot extrusion; techno-economic analysis included"],
            ["6–12 month preliminary demonstrable outcomes", "Optimized prototypes + biodegradation data by Month 12"],
            ["Completion by December 2027", "19-month schedule: June 2026 – December 2027"],
            ["Strengthen cotton competitive position", "Novel cotton-cellulose packaging displaces PE and wood-pulp alternatives"],
            ["Commercial applications with measurable market impact", "Target: flexible food wrap, courier pouches, hygiene product overwrap"],
        ],
        col_widths=[7, 9]
    )

    # --- 3. BACKGROUND ---
    add_heading(doc, "3. Background and Prior Work", 1)
    doc.add_paragraph(
        "Single-use plastics account for over 40% of global plastic production, with flexible "
        "polyethylene packaging representing one of the largest and fastest-growing segments. "
        "Regulatory bans on non-recyclable plastics (EU Single-Use Plastics Directive, India's "
        "PWM Amendment Rules 2022) and consumer preference for compostable materials are "
        "accelerating demand for bio-based alternatives."
    )
    doc.add_paragraph(
        "Cotton cellulose offers inherent biodegradability, renewability, and established supply "
        "chain infrastructure. Cotton linter pulp—a byproduct of cotton ginning—contains >95% "
        "α-cellulose and is currently underutilized relative to wood pulp in specialty packaging "
        "applications. Cotton Incorporated's mission to expand cotton utilization into emerging "
        "bio-based material markets aligns directly with this proposal."
    )
    add_heading(doc, "3.1 Prior Work in Our Laboratory", 2)
    doc.add_paragraph(
        "Our research group has successfully developed biodegradable film prototypes using the "
        "following baseline formulation platform:"
    )
    add_table(doc,
        ["Component", "Function", "Typical Proportion (w/w)"],
        [
            ["Cornstarch (maize starch)", "Primary polysaccharide matrix", "60–70%"],
            ["Pectin (citrus/apple derived)", "Natural binder and film-forming agent", "5–10%"],
            ["Glycerol", "Plasticizer – improves flexibility and elongation", "15–20%"],
            ["Water", "Solvent and dispersion medium", "Q.S."],
            ["Acetic acid (vinegar, 4–5%)", "pH modifier; promotes pectin gelation and starch gelatinization", "2–5%"],
        ],
        col_widths=[4, 6, 4]
    )
    doc.add_paragraph(
        "Preliminary results from our laboratory demonstrate that this polysaccharide-glycerin-acid "
        "system produces flexible, translucent films that disintegrate within 4–8 weeks under "
        "ambient soil burial conditions. The proposed project will adapt this proven processing "
        "platform by substituting cotton linter pulp (and optionally CNCs) for cornstarch, "
        "systematically optimizing formulation variables to achieve PE-comparable performance."
    )

    # --- 4. PROBLEM STATEMENT ---
    add_heading(doc, "4. Problem Statement and Research Gap", 1)
    doc.add_paragraph(
        "Despite cotton's natural biodegradability, cotton-derived films have not been systematically "
        "developed as PE replacements in flexible packaging. Existing bioplastic research focuses "
        "predominantly on PLA, PHA, and starch blends that do not utilize cotton feedstock. "
        "Critical gaps include:"
    )
    for gap in [
        "No published whole-formulation studies combining cotton linter pulp, pectin, and bio-plasticizers for cast film applications",
        "Limited barrier property data (WVTR, OTR) for cotton-cellulose films vs. LDPE benchmarks",
        "Absence of marine and compost biodegradation data for cotton-derived flexible films with finishing chemistries",
        "No techno-economic comparison of cotton linter pulp vs. wood pulp as packaging feedstock in the Indian context",
    ]:
        doc.add_paragraph(gap, style="List Bullet")

    # --- 5. OBJECTIVES ---
    add_heading(doc, "5. Project Objectives and Hypotheses", 1)
    add_heading(doc, "5.1 Primary Objective", 2)
    doc.add_paragraph(
        "Develop and optimize cotton linter pulp-based biodegradable films using a pectin-glycerin-acetic "
        "acid formulation platform, achieving mechanical and barrier properties suitable for flexible "
        "packaging applications, with validated environmental biodegradation."
    )
    add_heading(doc, "5.2 Specific Objectives", 2)
    objectives = [
        "Formulate cotton-cellulose films with ≥50% cotton-derived content by dry weight",
        "Achieve tensile strength ≥15 MPa and elongation at break ≥80% (approaching LDPE benchmarks)",
        "Characterize water vapor transmission rate (WVTR) and compare against LDPE control films",
        "Conduct ISO 14855-1 / ASTM D5338 compost biodegradation testing on optimized formulations",
        "Perform marine environment simulation studies (water column + sediment conditions)",
        "Complete preliminary techno-economic analysis of cotton linter pulp vs. cornstarch and wood pulp feedstocks",
    ]
    for i, obj in enumerate(objectives, 1):
        doc.add_paragraph(f"Objective {i}: {obj}", style="List Number")

    add_heading(doc, "5.3 Hypotheses", 2)
    doc.add_paragraph(
        "H1: Cotton linter pulp can partially or fully replace cornstarch in our established "
        "polysaccharide film formulation without loss of film integrity.\n"
        "H2: Pectin crosslinking mediated by acetic acid will improve cotton-cellulose film "
        "cohesion and reduce water sensitivity.\n"
        "H3: Glycerol plasticization at 15–20% (w/w) will yield elongation properties comparable "
        "to LDPE while maintaining adequate tensile strength.\n"
        "H4: Cotton-derived films will achieve ≥90% biodegradation within 180 days under "
        "industrial composting conditions."
    )

    # --- 6. METHODOLOGY ---
    add_heading(doc, "6. Methodology and Work Plan", 1)

    phases = [
        ("Phase 1: Feedstock Preparation and Formulation Screening (Months 1–4)",
         [
             "Procure bleached cotton linter pulp (Moirai Cotton Pulp Pvt. Ltd., Hisar) and food-grade pectin, glycerol, and acetic acid from verified vendors",
             "Prepare cotton linter pulp slurry (5–15% w/v) with controlled particle size reduction (homogenization / mild grinding)",
             "Optional: Extract cotton cellulose nanocrystals (CNCs) via sulfuric acid hydrolysis (64–98% H₂SO₄, 45°C, 45 min) for reinforcement trials",
             "Design 3×3×3 factorial experiment: cotton content (30%, 50%, 70%) × pectin (5%, 10%, 15%) × glycerol (10%, 15%, 20%)",
             "Solution casting on Teflon-coated plates; oven drying at 45–60°C; conditioning at 23°C / 50% RH for 48 h before testing",
             "Screen formulations by visual quality, handleability, and preliminary tensile testing",
         ]),
        ("Phase 2: Property Optimization and Benchmarking (Months 4–9)",
         [
             "Tensile testing on Universal Testing Machine (10 kN, Eqvimech) per ASTM D882",
             "Thickness measurement (digital micrometer, 10 random points per sample)",
             "Water vapor transmission rate (WVTR) per ASTM E96 gravimetric method (outsourced, J. K. Analytical Lab, Ahmedabad)",
             "Thermal analysis: DSC (glass transition, melting) and TGA (decomposition profile)",
             "Moisture absorption kinetics at 23°C / 50% and 23°C / 80% RH",
             "Direct comparison against LDPE control film and baseline cornstarch-pectin film",
             "Response surface methodology (RSM) to identify optimal formulation",
         ]),
        ("Phase 3: Environmental Validation (Months 9–14)",
         [
             "Compost biodegradation: ISO 14855-1 / ASTM D5338 respirometric CO₂ evolution (3 optimized formulations + controls)",
             "Disintegration testing per EN 13432 supplementary requirements",
             "Marine simulation: water column (oxic) and sediment (anoxic) microcosm studies per RFP Priority Area 1 guidance",
             "Assess byproduct formation, heavy metal residues, and ecotoxicity of compost end-product",
         ]),
        ("Phase 4: Techno-Economic Analysis and Reporting (Months 14–19)",
         [
             "Material cost modeling using actual vendor quotations (Indian supply chain)",
             "Energy and processing cost estimation for laboratory-to-pilot scale-up",
             "Life cycle comparison: cotton linter pulp film vs. LDPE vs. PLA (qualitative LCA)",
             "Prepare comprehensive final report, open-access publication, and Cotton Incorporated progress presentation",
         ]),
    ]
    for phase_title, steps in phases:
        add_heading(doc, phase_title, 2)
        for step in steps:
            doc.add_paragraph(step, style="List Bullet")

    add_heading(doc, "6.1 Formulation Protocol (Baseline – Adapted from Prior Work)", 2)
    add_table(doc,
        ["Step", "Procedure", "Parameters"],
        [
            ["1", "Dispersion", "Disperse cotton linter pulp in distilled water (90°C, 30 min stirring); add cornstarch if hybrid formulation"],
            ["2", "Gelatinization", "Cool to 60°C; add pectin pre-dissolved in warm water; stir 20 min"],
            ["3", "Plasticization", "Add glycerol; continue stirring 15 min at 50°C"],
            ["4", "Acidification", "Add acetic acid (or vinegar) dropwise to pH 4.0–4.5; stir 10 min"],
            ["5", "Casting", "Pour 40–60 mL onto Teflon plate (300×300 mm); spread uniformly with casting knife"],
            ["6", "Drying", "Air-dry 24 h + oven 50°C for 4 h; peel and condition before testing"],
        ],
        col_widths=[1.5, 6, 6]
    )

    # --- 7. TIMELINE ---
    add_heading(doc, "7. Timeline and Milestones", 1)
    add_table(doc,
        ["Month", "Activity", "Milestone / Deliverable"],
        [
            ["1–2", "Procurement, lab setup, feedstock characterization", "All materials received; baseline cornstarch film replicated"],
            ["3–4", "Formulation screening (27 combinations)", "Top 5 formulations identified"],
            ["5–6", "Mechanical & thermal characterization", "Tensile data report; DSC/TGA profiles"],
            ["7–8", "Barrier testing & RSM optimization", "Optimized formulation selected (≥50% cotton content)"],
            ["9–10", "Compost biodegradation testing initiated", "Samples submitted to Akshar Global / ERRL partner labs"],
            ["11–12", "Marine simulation + interim report", "6-month progress report to Cotton Incorporated; prototype samples"],
            ["13–14", "Biodegradation results analysis", "Biodegradation report (compost + marine)"],
            ["15–16", "Techno-economic analysis", "TEA report with vendor-based cost model"],
            ["17–18", "Manuscript preparation, scale-up recommendations", "Draft manuscript submitted"],
            ["19", "Final reporting", "Comprehensive final deliverables (December 2027)"],
        ],
        col_widths=[2, 5, 7]
    )

    # --- 8. OUTCOMES ---
    add_heading(doc, "8. Expected Outcomes and Deliverables", 1)
    add_table(doc,
        ["Deliverable", "Description", "Due"],
        [
            ["D1", "Optimized cotton-cellulose film formulation(s) with processing protocol", "Month 8"],
            ["D2", "Mechanical property dataset (tensile, elongation, modulus) vs. LDPE", "Month 8"],
            ["D3", "Barrier property report (WVTR, moisture uptake)", "Month 9"],
            ["D4", "Biodegradation assessment (compost + marine simulation)", "Month 14"],
            ["D5", "Techno-economic feasibility summary", "Month 17"],
            ["D6", "Peer-reviewed publication (open access)", "Month 18"],
            ["D7", "Comprehensive final report per Cotton Inc. protocols", "December 2027"],
            ["D8", "Prototype film samples (A4 size, 3 formulations) for Cotton Inc. evaluation", "Month 12"],
        ],
        col_widths=[1.5, 9, 2.5]
    )

    add_heading(doc, "8.1 Performance Targets", 2)
    add_table(doc,
        ["Parameter", "Target Value", "Test Method", "LDPE Benchmark (Reference)"],
        [
            ["Tensile strength", "≥ 15 MPa", "ASTM D882", "10–25 MPa"],
            ["Elongation at break", "≥ 80%", "ASTM D882", "100–500%"],
            ["Film thickness", "80–120 µm", "Digital micrometer", "50–100 µm"],
            ["WVTR", "< 10 g/m²/day", "ASTM E96", "3–8 g/m²/day"],
            ["Cotton-derived content", "≥ 50% dry weight", "Gravimetric", "N/A"],
            ["Compost biodegradation (180 days)", "≥ 90% CO₂ evolution", "ISO 14855-1", "N/A (non-biodegradable)"],
        ],
        col_widths=[4, 3, 3, 4]
    )

    # --- 9. COMMERCIALIZATION ---
    add_heading(doc, "9. Commercialization and Market Impact Potential", 1)
    doc.add_paragraph(
        "India produces approximately 5.5 million bales of cotton annually, generating an estimated "
        "1.5–2.0 million tonnes of ginning byproducts including linters. Currently, a significant "
        "fraction of cotton linter is exported as low-value pulp or discarded. This project creates "
        "a pathway to convert this feedstock into high-margin flexible packaging films."
    )
    doc.add_paragraph("Target market applications:")
    for m in [
        "Flexible food wraps and bread bags (replacing LDPE cling film)",
        "E-commerce courier pouches and mailers",
        "Hygiene product overwrap (diaper wraps, sanitary pad packaging)",
        "Agricultural mulch films (extended Phase 2 opportunity)",
    ]:
        doc.add_paragraph(m, style="List Bullet")
    doc.add_paragraph(
        "Based on vendor-quoted raw material costs (Section 11), the estimated material cost for "
        "optimized cotton-cellulose film is approximately ₹180–₹240 per kg at laboratory scale, "
        "with projected reduction to ₹90–₹130 per kg at pilot scale (500 kg/batch), competitive "
        "with imported PLA film (₹200–₹350/kg) while offering superior biodegradability credentials."
    )

    # --- 10. QUALIFICATIONS ---
    add_heading(doc, "10. Investigator Qualifications and Institutional Resources", 1)
    doc.add_paragraph(
        "[Dr. _________________________] is [Professor/Associate Professor] in the Department of "
        "[_________________________] at [Institution Name]. Research expertise includes biopolymer "
        "formulation, sustainable packaging materials, and polysaccharide chemistry. The PI's group "
        "has demonstrated capability in starch-pectin biodegradable film development and maintains "
        "laboratory facilities for polymer casting, mechanical testing, and thermal analysis."
    )
    doc.add_paragraph("Institutional resources available:")
    for r in [
        "Polymer processing and casting laboratory (50 m²)",
        "Access to FTIR, SEM (shared facility) for chemical and morphological characterization",
        "Existing humidity-controlled conditioning room",
        "Central instrumentation facility for advanced characterization as needed",
        "Institutional ethics and biosafety clearance for microbiological biodegradation studies",
    ]:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_page_break()

    # --- 11. DETAILED BUDGET ---
    add_heading(doc, "11. Detailed Budget with Vendor Quotations (Indian Rupees)", 1)
    doc.add_paragraph(
        "All unit prices listed below are sourced from publicly available vendor catalogues and "
        "IndiaMART/TradeIndia listings accessed in May 2026. Prices marked 'proforma' indicate "
        "institutional purchase estimates pending formal quotation at time of procurement. "
        "Goods & Services Tax (GST) @ 18% is applied to taxable items (equipment, consumables, "
        "chemicals, testing services). Manpower costs are exempt from GST."
    )

    budget_rows = []
    for item in BUDGET_ITEMS:
        sl, cat, desc, qty, unit, unit_price, vendor, loc, gst, line_excl = item
        budget_rows.append([
            str(sl), cat, desc[:60] + ("..." if len(desc) > 60 else ""),
            f"{qty} {unit}", fmt_inr(unit_price), vendor[:35],
            fmt_inr(line_excl), fmt_inr(gst), fmt_inr(line_excl + gst),
        ])

    total_excl = sum(i[9] for i in BUDGET_ITEMS)
    total_gst = sum(i[8] for i in BUDGET_ITEMS)
    grand_total = total_excl + total_gst

    add_table(doc,
        ["Sl.", "Category", "Item Description", "Qty", "Unit Price\n(excl. GST)", "Vendor", "Amount\n(excl. GST)", "GST\n(18%)", "Total\n(incl. GST)"],
        budget_rows,
        col_widths=[0.8, 1.8, 4.5, 1.2, 1.5, 2.5, 1.5, 1.2, 1.5]
    )

    add_heading(doc, "11.1 Item-wise Vendor and Location Details", 2)
    detail_rows = []
    for item in BUDGET_ITEMS:
        sl, cat, desc, qty, unit, unit_price, vendor, loc, gst, line_excl = item
        detail_rows.append([
            str(sl), desc, f"{qty} {unit}", fmt_inr(unit_price),
            vendor, loc, fmt_inr(line_excl + gst),
        ])
    add_table(doc,
        ["Sl.", "Full Item Description", "Quantity", "Unit Price\n(excl. GST)", "Vendor Name", "Vendor Location", "Line Total\n(incl. GST)"],
        detail_rows,
        col_widths=[0.7, 4.5, 1.2, 1.5, 2.8, 3.5, 1.5]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Total Project Cost (excluding GST): {fmt_inr(total_excl)}\n").bold = True
    p.add_run(f"Total GST (18% on taxable items): {fmt_inr(total_gst)}\n").bold = True
    p.add_run(f"GRAND TOTAL (INR): {fmt_inr(grand_total)}").bold = True
    r = p.runs[-1]
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x2E, 0x5D, 0x3A)

    # --- 12. BUDGET SUMMARY ---
    add_heading(doc, "12. Budget Summary by Category", 1)
    categories = {}
    for item in BUDGET_ITEMS:
        cat = item[1]
        categories[cat] = categories.get(cat, 0) + item[9] + item[8]
    summary_rows = [[cat, fmt_inr(amt), f"{amt/grand_total*100:.1f}%"] for cat, amt in categories.items()]
    summary_rows.append(["GRAND TOTAL", fmt_inr(grand_total), "100.0%"])
    add_table(doc, ["Budget Head", "Amount (INR, incl. GST)", "% of Total"], summary_rows, col_widths=[6, 5, 3])

    # --- 13. REFERENCES ---
    add_heading(doc, "13. References", 1)
    refs = [
        "ASTM D882-18. Standard Test Method for Tensile Properties of Thin Plastic Sheeting.",
        "ASTM D6400-21. Standard Specification for Labeling of Plastics Designed to be Aerobically Composted in Municipal or Industrial Facilities.",
        "ASTM E96/E.g. Standard Test Methods for Water Vapor Transmission of Materials.",
        "ISO 14855-1:2012. Determination of the ultimate aerobic biodegradability of plastic materials under controlled composting conditions.",
        "Cotton Incorporated (2027). Request for Proposals – Product Development and Implementation Division.",
        "RFP Priority Area 3: Biodegradable Alternatives to Single-Use Plastics.",
        "Thakur, R. et al. (2019). Progress in green polymer composites. Composites Part B, 176, 107217.",
        "Utracki, F.M. & Wilkie, C.A. (2012). Biopolymer Blends. Wiley-VCH.",
        "UGC (2025–26). Junior Research Fellowship Stipend Rates. ugc.gov.in.",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"[{i}] {ref}")

    doc.add_page_break()

    # --- APPENDIX A ---
    add_heading(doc, "Appendix A – Vendor Quotation Summary Sheet", 1)
    doc.add_paragraph(
        "The following table consolidates vendor contact details and quoted unit prices for all "
        "major procurement items. Formal purchase orders will be placed after Cotton Incorporated "
        "approval and institutional procurement clearance."
    )
    # Vendor summary from verified suppliers
    vendor_summary = [
        ["Moirai Cotton Pulp Private Limited", "Salemgarh, Hisar, Haryana – 125052", "Bleached cotton linter pulp", "₹130.00/kg (excl. GST)", "www.indiamart.com / moirai.in", "+91-8043886354"],
        ["Amit Hydrocolloids", "121, Raja Industrial Estate, PK Road, Mumbai – 400080", "Pectin powder, food/pharma grade", "₹1,098.00/kg (excl. GST)", "www.amithydrocolloids.com", "+91-8045805088"],
        ["Otto Kemi (OTTO Chemie Pvt. Ltd.)", "Mumbai, Maharashtra", "Glycerol GR 99%+ AR", "₹7,009.00/5 L (excl. GST)", "www.ottokemi.com", "Via website enquiry"],
        ["Triveni Chemicals India", "Vapi, Gujarat", "Maize starch food grade", "₹35.00/kg (excl. GST)", "www.trivenichemicalsindia.com", "Via website enquiry"],
        ["Alpha Chemika", "Andheri West, Mumbai – 400058", "Acetic acid glacial AR", "₹350.00/500 mL; ₹1,220.00/2.5 L", "www.alphachemika.co", "+91-2226317055"],
        ["Labitems India", "Online – PAN India", "Sulphuric acid AR 2.5 L", "₹1,049.50/2.5 L (listed price)", "www.labitems.co.in", "Online order"],
        ["HiMedia Laboratories Pvt. Ltd.", "MIDC Wagle Estate, Thane – 400604", "Sodium hydroxide pellets AR", "₹285.00/kg (catalogue estimate)", "www.himedialabs.com", "+91-22-6116 9797"],
        ["Eqvimech Private Limited", "Jekegram, Thane – 400606", "UTM 10 kN electronic", "₹4,70,000.00/unit (excl. GST)", "www.indiamart.com/eqvimech", "+91-7942558462"],
        ["Labindia Analytical Instruments", "Mumbai, Maharashtra", "Hotplate stirrer; pH meter", "₹12,500/unit; ₹6,500/unit", "www.labindia.com", "Via dealer network"],
        ["Royal Scientific / EIE Instruments", "Ahmedabad / Delhi", "Vacuum oven 25 L; casting plates", "₹78,000; ₹8,500", "IndiaMART verified listings", "Via IndiaMART"],
        ["Akshar Global Exports", "Kolkata, West Bengal", "Biodegradation & compost testing", "₹35,000/sample (ISO 14855); range ₹25,000–50,000", "www.aksharglobalexports.com", "Via website enquiry"],
        ["ERRL (Environmental Research & Remediation Lab)", "India", "Marine biodegradation simulation", "₹55,000/study (project quote)", "errl.co.in", "Via website enquiry"],
        ["J. K. Analytical Laboratory & Research Centre", "Ahmedabad, Gujarat", "DSC/TGA; WVTR testing", "₹2,200/sample; ₹4,500/sample", "TradeIndia listing", "Via TradeIndia"],
    ]
    add_table(doc,
        ["Vendor Name", "Address", "Product / Service", "Quoted Price (INR)", "Website", "Contact"],
        vendor_summary,
        col_widths=[3, 3.5, 3, 2.5, 2.5, 2]
    )

    # --- APPENDIX B ---
    add_heading(doc, "Appendix B – Pre-Proposal Submission Checklist", 1)
    checklist = [
        ("Target research category specified", "Priority Area 3a – Cotton-derived biodegradable films"),
        ("Pre-proposal format", "Formal research document (this document)"),
        ("Concept included", "Section 1 – Executive Summary"),
        ("Methodology included", "Section 6 – Methodology and Work Plan"),
        ("Budget requirements included", "Section 11 – Detailed Budget (INR with vendor details)"),
        ("Fewer than 3 pre-proposals per investigator", "This is pre-proposal #1 of maximum 3"),
        ("Submission deadline", "Sunday, 31 May 2026"),
        ("Submission email", "TCR-Proposals@cottoninc.com"),
        ("Contact person", "Renuka Dhandapani, Manager, Textile Chemistry Research"),
        ("Contact phone", "919.678.2448"),
        ("All research complete by", "December 2027"),
    ]
    add_table(doc, ["Requirement", "Status / Details"], checklist, col_widths=[6, 10])

  # --- SIGNATURE BLOCK ---
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph("Principal Investigator Signature: _________________________    Date: _______________")
    sig = doc.add_paragraph("Name: [Dr. _________________________]")
    sig = doc.add_paragraph("Head of Department Endorsement: _________________________    Date: _______________")

    doc.save(OUTPUT)
    return OUTPUT, grand_total, total_excl, total_gst


if __name__ == "__main__":
    path, grand, excl, gst = build_document()
    print(f"Generated: {path}")
    print(f"Total excl GST: INR {excl:,.2f}")
    print(f"Total GST: INR {gst:,.2f}")
    print(f"Grand Total: INR {grand:,.2f}")
