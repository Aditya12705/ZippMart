"""Concise 12-month Cotton Incorporated pre-proposal with justified, computed budget."""

from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = r"d:\Checkout System\proposals\Cotton_Incorporated_RFP_2027_PreProposal_Biodegradable_Cotton_Films.docx"
GST = 0.18


def shade(cell, color="D9EAD3"):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(el)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x2E, 0x5D, 0x3A)
        r.font.size = Pt(12 if level > 1 else 13)


def body(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        shade(t.rows[0].cells[i])
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
            for p in t.rows[ri].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    if widths:
        for ci, w in enumerate(widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    return t


def fmt(n):
    return f"₹{n:,.2f}"


def line_total(qty, rate, taxable=True):
    ex = round(qty * rate, 2)
    gst_amt = round(ex * GST, 2) if taxable else 0.0
    return ex, gst_amt, round(ex + gst_amt, 2)


# Each row: sl, head, qty, unit, rate, taxable, vendor, location, justification
RAW = [
    (
        "1",
        "Cotton linters (bleached / 1st cut)",
        165,
        "kg",
        25.00,
        True,
        "Vicram Yarns And Fibers",
        "Uttrahalli, Bengaluru 560061",
        "165 kg for ~15 casting campaigns (~10 kg each), linter grinding trials, and reserve for moisture equilibration repeats.",
    ),
    (
        "2",
        "Food-grade maize starch (baseline + hybrid films)",
        100,
        "kg",
        39.00,
        True,
        "Well Thought Chemicals",
        "Peenya, Bengaluru 560058",
        "100 kg for starch control films and starch–cotton blends through optimisation (Months 1–9).",
    ),
    (
        "3",
        "Pectin powder (food grade)",
        18,
        "kg",
        1098.00,
        True,
        "Amit Hydrocolloids (via Bengaluru distributor quote)",
        "Mumbai supply to Bengaluru",
        "18 kg: binder at 5–12% in repeated 2–3 kg batch scales across the formulation matrix.",
    ),
    (
        "4",
        "Glycerine (>99%, lab / industrial grade)",
        28,
        "kg",
        250.00,
        True,
        "Paxal Chemical Industry Pvt. Ltd.",
        "Sheshadripuram, Bengaluru 560020",
        "28 kg plasticizer for 15–20% loading trials on all active formulations.",
    ),
    (
        "5",
        "Glacial acetic acid",
        22,
        "kg",
        100.00,
        True,
        "Miracle Ingredients LLP",
        "Pantharapalya, Bengaluru 560039",
        "22 kg for pH control and vinegar-equivalent acidification in routine casting.",
    ),
    (
        "6",
        "Food-grade vinegar, caustic soda flakes, distilled water supplements",
        1,
        "lot",
        22400.00,
        True,
        "Well Thought Chemicals; local FSSAI supplier",
        "Bengaluru",
        "Minor reagents and repeat purchase of vinegar for pilot batches (proforma lot).",
    ),
]

CONSUMABLES = [
    (
        "7",
        "Borosilicate beakers, measuring cylinders, spatulas",
        1,
        "set",
        36400.00,
        True,
        "Nandini Marketing Company / Bharat Scientific World",
        "Bengaluru",
        "Dedicated glassware for parallel casting lanes (2–3 formulations/week in active phase).",
    ),
    (
        "8",
        "Teflon casting plates, levelling knives, release sheets",
        1,
        "set",
        22800.00,
        True,
        "Bharat Scientific World",
        "Bengaluru",
        "Non-stick casting surface for uniform film thickness (target 80–120 µm).",
    ),
    (
        "9",
        "Filter paper, aluminium foil, HDPE bags, desiccant, lab labels",
        1,
        "lot",
        18650.00,
        True,
        "Local scientific consumables dealer",
        "Bengaluru",
        "Conditioning and storage of films before FTIR, tensile, and biodegradation submission.",
    ),
]

EQUIPMENT = [
    (
        "10",
        "Magnetic hotplate stirrer with digital temperature control",
        2,
        "unit",
        12850.00,
        True,
        "Labindia Analytical Instruments (authorised dealer)",
        "Bengaluru",
        "Two units so baseline and cotton formulations can run in parallel during Months 2–8.",
    ),
    (
        "11",
        "Vacuum oven (25 L) for controlled drying",
        1,
        "unit",
        89500.00,
        True,
        "Regional lab equipment supplier (IndiaMART quote, Bengaluru delivery)",
        "Bengaluru",
        "Stable drying after casting; department oven is shared and not always available.",
    ),
    (
        "12",
        "Digital thickness gauge (0–10 mm)",
        1,
        "unit",
        9650.00,
        True,
        "Authorised Mitutoyo / Insize dealer",
        "Bengaluru",
        "Ten-point thickness per film for tensile and WVTR sample preparation.",
    ),
]

TESTING = [
    (
        "13",
        "FTIR — formulation fingerprinting & cotton–starch comparison",
        20,
        "sample",
        4500.00,
        True,
        "Dextrose Technologies Pvt. Ltd.",
        "Kengeri, Bengaluru 560060",
        "20 spectra: 8 variants × 2 prep + 4 repeat checks after formulation lock-in.",
    ),
    (
        "14",
        "DSC / TGA — thermal transitions & mass loss",
        22,
        "sample",
        3850.00,
        True,
        "Dextrose Technologies Pvt. Ltd.",
        "Kengeri, Bengaluru 560060",
        "18 runs on short-listed films plus starch control and blank pans.",
    ),
    (
        "15",
        "Tensile properties (ASTM D882) — 3 formulations, 5 specimens each",
        1,
        "study",
        168400.00,
        True,
        "Dextrose Technologies Pvt. Ltd. (UTM service quote)",
        "Kengeri, Bengaluru 560060",
        "One outsourced campaign; 15 dog-bone specimens cut from cast films (quoted as package).",
    ),
    (
        "16",
        "WVTR (ASTM E96 gravimetric) & moisture uptake panel",
        8,
        "sample",
        4650.00,
        True,
        "Nanowatts Technologies",
        "RMV 2nd Stage, Bengaluru 560094",
        "8 film sets: 3 leads + control, duplicate conditioning at 50% and 65% RH.",
    ),
    (
        "17",
        "Compost biodegradation (ISO 14855-1 / ASTM D5338)",
        3,
        "formulation",
        98500.00,
        True,
        "Greenlink Laboratory",
        "Bengaluru",
        "3 lead films; full respirometric panel (quoted range ₹90,000–₹1,00,000 per material).",
    ),
    (
        "18",
        "Disintegration / soil-burial support & compost extract screening",
        1,
        "panel",
        52800.00,
        True,
        "Greenlink Laboratory",
        "Bengaluru",
        "Supplementary checks aligned with RFP interest in realistic end-of-life behaviour.",
    ),
    (
        "19",
        "SEM — surface morphology of film fracture (6 specimens)",
        6,
        "sample",
        5800.00,
        True,
        "Dextrose Technologies Pvt. Ltd.",
        "Kengeri, Bengaluru 560060",
        "6 micrographs on lead films after tensile failure (outsourced imaging).",
    ),
]

MANPOWER = [
    (
        "20",
        "Project Research Fellow",
        12,
        "month",
        36000.00,
        False,
        "Host institution",
        "—",
        "Full-time bench work: casting, conditioning, sample prep, lab coordination (12 × ₹36,000; UGC JRF–aligned rate).",
    ),
    (
        "21",
        "Technical Assistant",
        12,
        "month",
        13200.00,
        False,
        "Host institution",
        "—",
        "Weighing, drying, cleaning, sample logging, vendor/lab runs (12 × ₹13,200).",
    ),
    (
        "22",
        "PI research allowance",
        12,
        "month",
        10800.00,
        False,
        "Host institution",
        "—",
        "Supervision, data review, report writing (12 × ₹10,800).",
    ),
]

OTHER = [
    (
        "23",
        "Domestic travel & sample logistics (Bengaluru vendors / testing labs)",
        9,
        "trip",
        6350.00,
        False,
        "—",
        "Karnataka",
        "9 trips for procurement, sample handover to Dextrose/Greenlink/Nanowatts, and vendor follow-up.",
    ),
]


def build_budget_rows():
    rows = []
    taxable_base = 0.0
    for group in (RAW, CONSUMABLES, EQUIPMENT, TESTING, MANPOWER, OTHER):
        for sl, head, qty, unit, rate, taxable, vendor, loc, just in group:
            ex, gst_amt, incl = line_total(qty, rate, taxable)
            if taxable:
                taxable_base += ex
            rows.append(
                {
                    "sl": sl,
                    "head": head,
                    "basis": f"{qty:g} {unit} × {fmt(rate)}/{unit}" if unit != "lot" and unit != "study" and unit != "panel" and unit != "formulation" and unit != "set" else (
                        f"{qty:g} {unit} @ {fmt(rate)}" if qty != 1 else f"1 {unit} @ {fmt(rate)}"
                    ),
                    "vendor": vendor,
                    "loc": loc,
                    "ex": ex,
                    "gst": gst_amt,
                    "incl": incl,
                    "just": just,
                    "taxable": taxable,
                }
            )
    cont_ex = round(taxable_base * 0.05, 2)
    cont_gst = round(cont_ex * GST, 2)
    cont_incl = round(cont_ex + cont_gst, 2)
    rows.append(
        {
            "sl": "24",
            "head": "Contingency",
            "basis": f"5% of taxable direct cost ({fmt(taxable_base)} excl. GST)",
            "vendor": "—",
            "loc": "—",
            "ex": cont_ex,
            "gst": cont_gst,
            "incl": cont_incl,
            "just": "Price escalation, extra casting batches, or repeat of 2–3 failed FTIR/tensile specimens.",
            "taxable": True,
        }
    )
    grand = round(sum(r["incl"] for r in rows), 2)
    return rows, grand, taxable_base


def build():
    budget_rows, grand_total, _ = build_budget_rows()

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRE-PROPOSAL — COTTON INCORPORATED RFP 2027")
    r.bold = True
    r.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Cotton-Linter Reinforced Biodegradable Films from a\n"
        "Starch–Pectin–Glycerine Formulation Platform"
    )
    r.font.size = Pt(12)

    for line in [
        "[Dr. _________________]  |  [Department]  |  [Institution, India]",
        "Project period: 12 months (completion by December 2027)",
        f"Date: {date.today().strftime('%d %B %Y')}",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(10)

    doc.add_paragraph()

    heading(doc, "1. Work Already Carried Out in Our Laboratory")
    body(
        doc,
        "We have prepared biodegradable sheets using cornstarch, pectin, glycerine, water, and "
        "vinegar (dilute acetic acid). Films were cast from a heated starch–pectin slurry with "
        "glycerine as plasticizer and mild acid to set the network. The sheets were flexible, "
        "semi-transparent, and showed visible breakdown within about 4–8 weeks in simple soil-burial "
        "checks. The process is repeatable at bench scale.",
    )
    body(
        doc,
        "That work is the starting point. This project adds cotton linters to the same route and "
        "adds the testing needed to judge whether the films are viable for flexible packaging.",
    )

    heading(doc, "2. Fit with Cotton Incorporated Priority Areas")
    table(
        doc,
        ["RFP priority area", "How our project addresses it"],
        [
            [
                "3a — Cotton-derived films / wraps as alternatives to polyethylene",
                "Cotton linters are built into the film matrix; end use is flexible wrap/pouch-type packaging.",
            ],
            [
                "3 — Biodegradable alternatives to single-use plastics",
                "All major inputs are bio-based; no petroleum polymer binder.",
            ],
            [
                "5 — Valorization of cotton processing byproducts",
                "Cotton linters (ginning byproduct) are the main cotton input, sourced in Bengaluru.",
            ],
            [
                "1b — Realistic biodegradation assessment (supporting)",
                "Compost respirometry and soil/disintegration screening are included in Months 10–12.",
            ],
        ],
        widths=[5.2, 10.3],
    )

    heading(doc, "3. Objectives (12 Months)")
    for o in [
        "Integrate cotton linters (10–40% dry basis) into the starch–pectin–glycerine process.",
        "Select 2–3 formulations with acceptable flexibility, thickness, and handling.",
        "Obtain FTIR, DSC/TGA, tensile (ASTM D882), and WVTR data against a starch-only control.",
        "Run compost biodegradation on the lead film(s) and document results by December 2027.",
    ]:
        p = doc.add_paragraph(o, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)

    heading(doc, "4. Work Plan (Quarterly)", 2)
    table(
        doc,
        ["Period", "Activities", "Deliverable"],
        [
            ["Months 1–3", "Materials procurement; baseline films; cotton loading 10–30%.", "Baseline + cotton films."],
            ["Months 4–6", "Shortlist 3 formulations; FTIR; start DSC/TGA.", "FTIR/TGA report."],
            ["Months 7–9", "Tensile (D882), WVTR; adjust plasticizer/pectin if needed.", "Mechanical/barrier data."],
            ["Months 10–12", "Compost test; soil checks; final report and film samples.", "Biodegradation summary."],
        ],
        widths=[2.0, 8.5, 5.0],
    )

    heading(doc, "5. Budget and Justification (12 Months)")
    body(
        doc,
        "The budget is built from (i) quantities needed for repeated casting and optimisation over "
        "12 months, (ii) vendor-listed unit rates for Bengaluru suppliers where available, and "
        "(iii) quoted per-sample or per-study fees from Bengaluru analytical laboratories. GST at "
        "18% is applied to taxable purchases and services. Staff costs follow institutional project "
        "fellowship norms. Contingency is 5% of taxable direct costs only. The project total is the "
        "sum of the line items below (no lump-sum rounding to a target figure).",
    )

    table(
        doc,
        ["Sl.", "Item", "Quantity basis", "Amount (incl. GST)", "Justification"],
        [
            [
                r["sl"],
                r["head"],
                r["basis"],
                fmt(r["incl"]),
                r["just"][:120] + ("…" if len(r["just"]) > 120 else ""),
            ]
            for r in budget_rows
        ]
        + [["", "Total estimated project cost", "", fmt(grand_total), "Sum of lines 1–24"]],
        widths=[0.6, 3.2, 2.8, 2.2, 6.7],
    )

    heading(doc, "5.1 Summary by Category", 2)
    cats = {}
    for r in budget_rows:
        if r["sl"] == "24":
            key = "Contingency"
        elif int(r["sl"]) <= 6:
            key = "Raw materials"
        elif int(r["sl"]) <= 9:
            key = "Consumables"
        elif int(r["sl"]) <= 12:
            key = "Equipment"
        elif int(r["sl"]) <= 19:
            key = "Testing & analysis"
        elif int(r["sl"]) <= 22:
            key = "Manpower"
        else:
            key = "Travel & logistics"
        cats[key] = cats.get(key, 0) + r["incl"]
    summary = [[k, fmt(v), f"{v / grand_total * 100:.1f}%"] for k, v in cats.items()]
    summary.append(["Total", fmt(grand_total), "100%"])
    table(doc, ["Category", "Amount", "Share"], summary, widths=[5, 4, 3])

    heading(doc, "5.2 Vendors (Bengaluru / Karnataka)", 2)
    vrows = [
        [r["head"][:42], r["vendor"][:30], r["loc"][:35], fmt(r["incl"])]
        for r in budget_rows
        if r["vendor"] != "—"
    ]
    table(doc, ["Item", "Vendor", "Location", "Allocation"], vrows, widths=[4.2, 3.2, 4.0, 2.1])

    heading(doc, "6. Expected Outputs and Applications", 2)
    body(
        doc,
        "By December 2027: documented film recipe(s); FTIR/thermal/mechanical/WVTR dataset; compost "
        "biodegradation results; and prototype films for Cotton Incorporated review.",
    )
    body(
        doc,
        "If mechanical strength and moisture resistance meet targets, the same film platform could "
        "support short-life flexible packaging where polyethylene is used today but composting or "
        "soil return is preferred:",
    )
    for app in [
        "Fresh produce overwrap and tray liners (e.g. cucumbers, herbs, berries) in retail and farm-gate packing.",
        "Bread, bakery, and dry-snack overwrap where a breathable, compostable sheet replaces LDPE cling.",
        "Garment and textile shipment bags for cotton apparel brands—cotton-linter content aligns with fibre sourcing narratives.",
        "Protective wrap for cotton lint, yarn cones, or small textile rolls in mills and warehouses (byproduct-to-packaging loop).",
        "Compostable mailers and pouch inserts for e-commerce parcels with low liquid exposure.",
        "Horticulture: seedling collars, pot liners, or mulch sheets that can remain in soil at end of season.",
        "Floral and gift overwrap for events and florists seeking plastic-free presentation.",
    ]:
        p = doc.add_paragraph(app, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)
    body(
        doc,
        "These are illustrative end uses; final adoption would depend on tensile/WVTR data from this "
        "project and any food-contact or brand-specific certification required by the customer.",
    )

    doc.add_paragraph()
    doc.add_paragraph("PI signature: _________________________    Date: _________")
    doc.add_paragraph("HOD endorsement: _________________________    Date: _________")

    doc.save(OUTPUT)
    return OUTPUT, grand_total


if __name__ == "__main__":
    path, total = build()
    print(f"Generated: {path}")
    print(f"Estimated total: INR {total:,.2f}")
