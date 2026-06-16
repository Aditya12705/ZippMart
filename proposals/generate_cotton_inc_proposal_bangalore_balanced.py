"""Generate a balanced Bangalore-based Cotton Incorporated pre-proposal."""

from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = r"d:\Checkout System\proposals\Cotton_Incorporated_RFP_2027_PreProposal_Biodegradable_Cotton_Films.docx"
GST_RATE = 0.18


def set_shading(cell, color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x5D, 0x3A)
    return heading


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
        set_shading(table.rows[0].cells[i], "D9EAD3")
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = str(value)
            for p in table.rows[r].cells[c].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for c, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[c].width = Cm(width)
    return table


def fmt_inr(value):
    return f"₹{value:,.2f}"


PROCURED_ITEMS = [
    {
        "sl": 1,
        "category": "Raw material",
        "item": "Cotton linters for cotton-derived reinforcement and filler trials",
        "qty": 50,
        "unit": "kg",
        "unit_price": 25.00,
        "vendor": "Vicram Yarns And Fibers",
        "address": "303, 2nd Floor, Siri Aroma Apartments, 90 Kempanna Layout, Gowdanapalya, Uttrahalli, Bengaluru, Karnataka 560061",
        "basis": "TradeIndia listing: Cotton Linters at ₹25/kg in Bengaluru",
    },
    {
        "sl": 2,
        "category": "Raw material",
        "item": "Food-grade maize starch for baseline and hybrid film formulations",
        "qty": 50,
        "unit": "kg",
        "unit_price": 39.00,
        "vendor": "Well Thought Chemicals",
        "address": "355, 9th Cross, 4th Main Road, 4th Phase, Peenya Industrial Area, Near NTTF Circle, Bengaluru, Karnataka 560058",
        "basis": "IndiaMART listing: White Maize Starch Powder at ₹39/kg in Bengaluru",
    },
    {
        "sl": 3,
        "category": "Raw material",
        "item": "Glycerine (>99%) for plasticization and flexibility control",
        "qty": 10,
        "unit": "kg",
        "unit_price": 250.00,
        "vendor": "Paxal Chemical Industry Private Limited",
        "address": "Municipal No. 6, Railway Parallel Road, Nehru Nagar, Sheshadripuram, Bengaluru, Karnataka 560020",
        "basis": "IndiaMART listing: Glycerine for laboratory use at ₹250/kg",
    },
    {
        "sl": 4,
        "category": "Raw material",
        "item": "Glacial acetic acid for controlled acidification of the formulation",
        "qty": 15,
        "unit": "kg",
        "unit_price": 100.00,
        "vendor": "Miracle Ingredients LLP",
        "address": "No. 4, 2nd Cross, Sri Ranganagara Outer Ring Road, Pantharapalya, 100 Ft Ring Road, Bengaluru, Karnataka 560039",
        "basis": "Bengaluru listing: Glacial Acetic Acid approx. ₹100/kg",
    },
    {
        "sl": 5,
        "category": "Supporting chemical",
        "item": "Caustic soda flakes for cleaning, dispersion correction, and limited pH adjustment",
        "qty": 10,
        "unit": "kg",
        "unit_price": 54.00,
        "vendor": "Well Thought Chemicals",
        "address": "355, 9th Cross, 4th Main Road, 4th Phase, Peenya Industrial Area, Near NTTF Circle, Bengaluru, Karnataka 560058",
        "basis": "IndiaMART listing: Caustic Soda Flakes at ₹54/kg in Bengaluru",
    },
    {
        "sl": 6,
        "category": "Labware",
        "item": "Borosilicate beakers dedicated for repeated casting and formulation batches",
        "qty": 4,
        "unit": "piece",
        "unit_price": 500.00,
        "vendor": "Nandini Marketing Company",
        "address": "No. 174, Nandini Mansion, 3rd Main, 5th Cross, BHEL Layout, Kenchenhalli, Rajarajeshwari Nagar, Bengaluru, Karnataka 560098",
        "basis": "Supplier listing: BORO+ Borosilicate Beaker at ₹500/piece",
    },
]


for item in PROCURED_ITEMS:
    item["amount_ex_gst"] = round(item["qty"] * item["unit_price"], 2)
    item["gst"] = round(item["amount_ex_gst"] * GST_RATE, 2)
    item["amount_incl_gst"] = round(item["amount_ex_gst"] + item["gst"], 2)

MATERIAL_SUBTOTAL = round(sum(item["amount_ex_gst"] for item in PROCURED_ITEMS), 2)
GST_TOTAL = round(sum(item["gst"] for item in PROCURED_ITEMS), 2)
STUDENT_ASSISTANT = 48000.00  # 6 months x ₹8,000/month
CONTINGENCY = round(MATERIAL_SUBTOTAL * 0.05, 2)
GRAND_TOTAL = round(MATERIAL_SUBTOTAL + GST_TOTAL + STUDENT_ASSISTANT + CONTINGENCY, 2)


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRE-PROPOSAL FOR RESEARCH FUNDING")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x5D, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Request for Proposals 2027\nProduct Development and Implementation Division\nCotton Incorporated")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Development of Cotton-Linter Reinforced Biodegradable Films\n"
        "from a Starch-Pectin-Glycerine Formulation Platform"
    )
    run.italic = True
    run.font.size = Pt(13)

    doc.add_paragraph()
    cover = [
        ("Principal Investigator:", "[Dr. _________________________]"),
        ("Department:", "[Department of _________________________]"),
        ("Institution:", "[University / Institute Name, India]"),
        ("Target Category:", "Priority Area 3a - Cotton-derived biodegradable alternatives to polyethylene"),
        ("Project Duration:", "6 months (preliminary proof-of-concept stage)"),
        ("Budget Requested:", fmt_inr(GRAND_TOTAL)),
        ("Procurement:", "Quoted purchase items proposed from Bengaluru-based vendors"),
        ("Document Date:", date.today().strftime("%d %B %Y")),
    ]
    for label, value in cover:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label} ").bold = True
        p.add_run(value)

    doc.add_page_break()

    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "This proposal seeks support for a six-month study on cotton-linter reinforced biodegradable "
        "films for single-use packaging applications. The work builds on our existing laboratory "
        "experience with starch-pectin-glycerine films and extends that formulation by introducing "
        "cotton linters as a cotton-derived reinforcing component. The study is designed to generate "
        "preliminary data on formulation stability, film formation, handling characteristics, and "
        "initial biodegradation behavior."
    )
    doc.add_paragraph(
        "The proposed work is focused on laboratory-scale formulation development and repeat trials. "
        "Support is requested for essential raw materials, limited laboratory support items, part-time "
        "student assistance for sustained experimental work, and contingency for routine project "
        "expenses. Existing institutional facilities will be used for preparation and basic evaluation."
    )

    add_heading(doc, "2. Project Rationale and Feasibility", 1)
    for point in [
        "Repeated formulation trials are necessary to identify a stable cotton-linter loading range and reproducible casting conditions.",
        "Routine bench work including mixing, casting, drying, weighing, sample conditioning, and biodegradation observation requires dedicated project time.",
        "The work can be completed with existing institutional equipment, while the requested budget is limited to essential consumables and minimum project support.",
        "Major capital purchases and certification-scale external testing are not proposed at this stage of the study.",
    ]:
        doc.add_paragraph(point, style="List Bullet")

    add_heading(doc, "3. Alignment with Cotton Incorporated RFP", 1)
    add_table(
        doc,
        ["RFP priority", "Project fit"],
        [
            [
                "Priority Area 3a - Cotton-derived films, wraps, or barrier coatings",
                "Cotton linters are used directly in biodegradable film development as a cotton-derived reinforcement component.",
            ],
            [
                "Priority Area 5 - Valorization of cotton processing byproducts",
                "Cotton linters, an underutilized byproduct, are converted into a value-added biodegradable material.",
            ],
            [
                "Targeted, near-term preliminary outcomes",
                "The project is intentionally limited to 6 months and focused on formulation optimization and initial performance evidence.",
            ],
        ],
        col_widths=[6, 10],
    )

    add_heading(doc, "4. Technical Background", 1)
    doc.add_paragraph(
        "Our laboratory has already made biodegradable plastic sheets using cornstarch, pectin, "
        "water, vinegar, and glycerine. That work established a reliable preparation route, but it "
        "did not yet incorporate cotton. The proposed study will retain the successful film-forming "
        "base and introduce cotton linters in increasing proportions so that the project remains "
        "grounded in prior experience while becoming directly relevant to Cotton Incorporated."
    )
    doc.add_paragraph(
        "Pectin required for the initial formulation work is already available in the laboratory from "
        "ongoing biodegradable film studies; therefore, no additional procurement is proposed under "
        "that head for the present study."
    )

    add_heading(doc, "5. Objectives", 1)
    objectives = [
        "To prepare cotton-linter reinforced biodegradable films using an established starch-pectin-glycerine formulation route.",
        "To evaluate whether cotton linters improve stiffness, handleability, and structural integrity of the films.",
        "To identify one or two optimized low-cost formulations suitable for a later full proposal.",
        "To generate preliminary data strong enough to justify larger-scale testing in the invited proposal stage.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Number")

    add_heading(doc, "6. Scope of Work for 6 Months", 1)
    phases = [
        ("Month 1", "Replication of prior starch-pectin formulation and baseline film preparation"),
        ("Month 2", "Cotton linter incorporation trials at multiple loading levels"),
        ("Month 3", "Optimization of mixing, drying, and casting conditions"),
        ("Month 4", "Repeat batches for the best-performing formulations"),
        ("Month 5", "Initial mechanical handling comparison and biodegradation observation"),
        ("Month 6", "Data consolidation and full-proposal preparation"),
    ]
    add_table(doc, ["Timeline", "Activity"], phases, col_widths=[3, 13])

    add_heading(doc, "7. Project Resources", 1)
    add_table(
        doc,
        ["Included in the proposal budget", "Available through institutional support"],
        [
            [
                "Core materials, limited supporting chemicals, dedicated glassware, one part-time student assistant, contingency",
                "Existing pectin stock, departmental instruments, hotplate/stirrer, oven, balance, routine glassware, faculty supervision",
            ],
            [
                "Repeated lab-scale film preparation",
                "Existing preparation and drying facilities",
            ],
            [
                "Sustained bench work over six months",
                "Advanced certification-scale external testing at a later stage, if required",
            ],
        ],
        col_widths=[8, 8],
    )

    add_heading(doc, "8. Detailed Budget", 1)
    budget_rows = []
    for item in PROCURED_ITEMS:
        budget_rows.append(
            [
                item["sl"],
                item["category"],
                item["item"],
                f'{item["qty"]} {item["unit"]}',
                fmt_inr(item["unit_price"]),
                item["vendor"],
                fmt_inr(item["amount_ex_gst"]),
                fmt_inr(item["gst"]),
                fmt_inr(item["amount_incl_gst"]),
            ]
        )
    add_table(
        doc,
        ["Sl.", "Category", "Item", "Qty", "Unit price", "Bengaluru vendor", "Amount ex GST", "GST", "Amount incl GST"],
        budget_rows,
        col_widths=[0.7, 1.7, 4.8, 1.2, 1.5, 3.0, 1.5, 1.2, 1.6],
    )

    add_heading(doc, "8.1 Project Support Costs", 2)
    support_rows = [
        ["Part-time student project assistant", "6 months x ₹8,000/month", fmt_inr(STUDENT_ASSISTANT)],
        ["Contingency", "5% of procured material cost (ex GST)", fmt_inr(CONTINGENCY)],
    ]
    add_table(doc, ["Budget head", "Basis", "Amount"], support_rows, col_widths=[5, 7, 3])

    p = doc.add_paragraph()
    p.add_run(f"Procured items subtotal (ex GST): {fmt_inr(MATERIAL_SUBTOTAL)}\n").bold = True
    p.add_run(f"GST on procured items: {fmt_inr(GST_TOTAL)}\n").bold = True
    p.add_run(f"Student assistant support: {fmt_inr(STUDENT_ASSISTANT)}\n").bold = True
    p.add_run(f"Contingency: {fmt_inr(CONTINGENCY)}\n").bold = True
    final_run = p.add_run(f"Grand Total Requested: {fmt_inr(GRAND_TOTAL)}")
    final_run.bold = True
    final_run.font.size = Pt(12)
    final_run.font.color.rgb = RGBColor(0x2E, 0x5D, 0x3A)

    add_heading(doc, "9. Vendor Basis for Priced Items", 1)
    vendor_rows = []
    for item in PROCURED_ITEMS:
        vendor_rows.append(
            [
                item["vendor"],
                item["address"],
                item["basis"],
            ]
        )
    add_table(doc, ["Vendor", "Address", "Price basis"], vendor_rows, col_widths=[3.5, 7.0, 5.0])

    add_heading(doc, "9.1 Bengaluru-Only Quotation and Location Summary", 2)
    location_rows = []
    for item in PROCURED_ITEMS:
        location_rows.append(
            [
                item["item"],
                item["vendor"],
                item["address"],
                fmt_inr(item["unit_price"]),
                f'{item["qty"]} {item["unit"]}',
                fmt_inr(item["amount_incl_gst"]),
            ]
        )
    add_table(
        doc,
        ["Item", "Quoted Bengaluru Vendor", "Location", "Quoted Unit Rate", "Required Quantity", "Allocated Total"],
        location_rows,
        col_widths=[4.0, 3.0, 5.0, 1.5, 1.6, 1.6],
    )

    add_heading(doc, "10. Expected Preliminary Outcomes", 1)
    outcomes = [
        "A cotton-linter reinforced biodegradable film formulation with repeatable preparation method.",
        "A direct comparison between starch-based film controls and cotton-linter reinforced variants.",
        "Initial evidence on whether cotton linters improve the practicality of the biodegradable film system.",
        "A strong basis for a fuller invited proposal with advanced characterization and standard testing.",
    ]
    for outcome in outcomes:
        doc.add_paragraph(outcome, style="List Bullet")

    add_heading(doc, "11. Conclusion", 1)
    doc.add_paragraph(
        "The proposed study will establish whether cotton linters can be integrated effectively into "
        "a biodegradable starch-pectin film system for single-use packaging applications. By combining "
        "an already workable laboratory formulation with a cotton-derived reinforcing component, the "
        "project offers a practical route to generate early technical evidence for cotton-based "
        "biodegradable materials and to support the next stage of product development."
    )

    doc.add_paragraph()
    doc.add_paragraph("Principal Investigator Signature: _________________________    Date: _______________")
    doc.add_paragraph("Name: [Dr. _________________________]")
    doc.add_paragraph("Head of Department Endorsement: _________________________    Date: _______________")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Generated: {path}")
    print(f"Grand total: INR {GRAND_TOTAL:,.2f}")
